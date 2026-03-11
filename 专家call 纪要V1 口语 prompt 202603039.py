from openai import OpenAI
import json
import time
import re
from datetime import timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from concurrent.futures import ThreadPoolExecutor, as_completed


def read_file_content(file_path):
    """统一读取txt和docx文件"""
    file_path = Path(file_path)
    
    if file_path.suffix.lower() == '.docx':
        doc = Document(file_path)
        lines = []
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text or lines:
                lines.append(text)
        return '\n'.join(lines)
    else:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()


class GeminiHandler:
    def __init__(self):
        self.client = OpenAI(
            api_key="",
            base_url=""
        )
    
    def request_gpt(self, system_prompt, user_prompt, output_type="text", temperature=0, max_retries=3):
        """通用API调用，带重试机制"""
        msg = [
            {"role": "system", "content": [{"type": "text", "text": system_prompt}]},
            {"role": "user", "content": [{"type": "text", "text": user_prompt}]}
        ]
        
        for attempt in range(max_retries):
            try:
                # 根据output_type动态构建参数
                params = {
                    "model": "gemini-2.5-pro-thinking",
                    "temperature": temperature,
                    "messages": msg,
                    "stream": True
                }
                
                # 只在需要JSON时才添加response_format参数
                if output_type == "json":
                    params["response_format"] = {"type": "json_object"}
                
                response = self.client.chat.completions.create(**params)
                
                full_response = ""
                for chunk in response:
                    if content := chunk.choices[0].delta.content:
                        full_response += content
                
                return {"code": 0, "message": full_response}
                
            except Exception as e:
                print(f"API调用失败（尝试{attempt+1}/{max_retries}）: {str(e)}")
                if attempt < max_retries - 1:
                    time.sleep(2)
                    continue
                return {"code": 1, "message": f"错误: {str(e)}"}


class TranscriptSegmenter:
    """步骤1: 分段"""
    def __init__(self, gemini_handler):
        self.gemini = gemini_handler
        self.segment_prompt = """
# 任务
你是一个对话分析引擎。你的任务是：
1.  **分割对话**: 将长篇对话逐字稿，按照逻辑话题分割成多个片段。
2.  **识别受访专家**: 在这次专家访谈中，找出核心的**受访专家**。TA是提供最多核心信息和专业见解的一方，通常扮演回答问题的角色，而不是主要提问的主持人。请分析对话并给出其ID。


# 分割标准
1.  **逻辑优先**: 分割的首要依据是**话题转换**。一个新片段的**第一句话**通常是开启新话题的**提问句**。当一个问题明显将讨论引向新方向时，分割点应在该问题**之前**，确保这个提问句成为新片段的开始。
2.  **时长限制**: 每个片段时长不应超过15分钟。若单一话题过长，可在其内部的子话题转换点进行切分。理想时长为5-10分钟。
3.  **完整覆盖 (MECE)**: 所有片段必须完整覆盖整个对话记录。一个片段的 `end_time` 是其最后一句话的时间戳；下一个片段的 `start_time` 必须是紧接着的下一句话的时间戳。
4.  **时间戳精确**: `start_time` 必须是片段第一句话的时间，`end_time` 必须是片段最后一句话的时间。


# 输入格式
纯文本，包含说话人、时间戳和内容。
```
[...文件头...]
文字记录:
说话人 A 00:01
[内容]
说话人 B 00:05
[内容]
...
```


# 输出格式
严格遵循以下 JSON 对象格式。**你的唯一输出必须是纯 JSON，不包含任何解释或 markdown 标记。**


```
{
  "expert_speaker_id": "你识别出的核心专家发言人ID (例如: '说话人 2')",
  "segments": [
    {
      "start_time": "HH:MM",
      "end_time": "HH:MM",
      "topic": "简短概括该片段的核心议题"
    }
  ]
}
```"""
        
    def segment_transcript(self, transcript_path):
        """调用API对transcript进行分段，并识别受访专家"""
        print("=" * 60)
        print("步骤 1/3: 正在对transcript进行分段...")
        print("=" * 60)
        
        transcript_content = read_file_content(transcript_path)
        
        result = self.gemini.request_gpt(
            system_prompt=self.segment_prompt,
            user_prompt=transcript_content,
            output_type="json",
            temperature=0
        )
        
        if result['code'] == 0:
            try:
                response_json = json.loads(result['message'])
                expert_speaker_id = response_json.get('expert_speaker_id', '')
                segments_json = response_json.get('segments', [])
                
                print(f"✓ 识别专家: {expert_speaker_id}")
                print(f"✓ 分段成功，共 {len(segments_json)} 个段落")
                return expert_speaker_id, segments_json
            except:
                print("✗ 分段结果解析失败")
                return None, None
        else:
            print("✗ 分段API调用失败")
            return None, None


class TranscriptMerger:
    """步骤2: 合并分段"""
    @staticmethod
    def parse_time(time_str: str) -> timedelta:
        """将 MM:SS 或 HH:MM:SS 格式转换为 timedelta"""
        parts = time_str.split(':')
        if len(parts) == 2:
            minutes, seconds = map(int, parts)
            return timedelta(minutes=minutes, seconds=seconds)
        elif len(parts) == 3:
            hours, minutes, seconds = map(int, parts)
            return timedelta(hours=hours, minutes=minutes, seconds=seconds)
        else:
            raise ValueError(f"不支持的时间格式: {time_str}")
    
    @staticmethod
    def load_transcript_with_format(transcript_path: str) -> list:
        """读取逐字稿，保留原始格式"""
        time_pattern = re.compile(r'^(?:说话人|发言人)\s*\d+\s+(\d{1,2}:\d{2}(?::\d{2})?)')
        entries = []
        current_time = None
        current_block = []
        
        content = read_file_content(transcript_path)
        lines = content.split('\n')
        
        for line in lines:
            line = line.rstrip()
            m = time_pattern.match(line)
            
            if m:
                if current_block:
                    entries.append((current_time, '\n'.join(current_block)))
                    current_block = []
                
                current_time = TranscriptMerger.parse_time(m.group(1))
                current_block.append(line)
            elif line.strip() and current_time is not None:
                current_block.append(line)
        
        if current_block:
            entries.append((current_time, '\n'.join(current_block)))
        
        return entries
    
    @staticmethod
    def merge_segments(segments_json, transcript_path):
        """根据JSON时间段合并transcript，返回内存数据"""
        print("\n" + "=" * 60)
        print("步骤 2/3: 正在合并分段...")
        print("=" * 60)
        
        entries = TranscriptMerger.load_transcript_with_format(transcript_path)
        merged_segments = []
        
        for i, seg in enumerate(segments_json, start=1):
            start = TranscriptMerger.parse_time(seg['start_time'])
            end = TranscriptMerger.parse_time(seg['end_time'])
            blocks = [text for (ts, text) in entries if start <= ts <= end]
            segment_text = '\n\n'.join(blocks)
            
            merged_segments.append({
                'topic': seg.get('topic', ''),
                'content': segment_text
            })
            
            print(f"段落 {i}: {seg.get('topic', '未命名')}")
            print(f"  时间范围: {seg['start_time']} - {seg['end_time']}")
            print(f"  内容长度: {len(segment_text)} 字符")
        
        print(f"✓ 合并完成，共 {len(merged_segments)} 个段落\n")
        
        return merged_segments


 
class MarkdownToWordConverter:
    """将Markdown转换为Word文档"""
    
    @staticmethod
    def set_font(run, font_name='微软雅黑', font_size=11, bold=False):
        """设置字体，同时处理中英文"""
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    @staticmethod
    def convert(md_content, output_path):
        """将MD内容转换为Word文档"""
        doc = Document()
        
        # 设置默认样式
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        lines = md_content.split('\n')
        i = 0
        
        # 列表项匹配正则：捕获缩进和内容
        list_pattern = re.compile(r'^(\s*)[-*]\s+(.*)$')
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # 跳过空行
            if not line:
                i += 1
                continue
            
            # 处理三级标题 ###
            if line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
                for run in p.runs:
                    MarkdownToWordConverter.set_font(run, '微软雅黑', 14, True)
                i += 1
                
            # 处理列表项（支持任意缩进层级）
            else:
                match = list_pattern.match(line)
                if match:
                    indent_spaces = len(match.group(1))
                    content = match.group(2)
                    
                    # 根据缩进空格数确定层级：0空格=层级0，2-3空格=层级1，4+空格=层级2
                    if indent_spaces == 0:
                        indent_level = 0
                    elif indent_spaces <= 3:
                        indent_level = 1
                    else:
                        indent_level = 2
                    
                    # 使用对应的列表样式
                    if indent_level == 0:
                        p = doc.add_paragraph(style='List Bullet')
                    elif indent_level == 1:
                        p = doc.add_paragraph(style='List Bullet 2')
                    else:
                        p = doc.add_paragraph(style='List Bullet 3')
                    
                    MarkdownToWordConverter._add_formatted_text(p, content)
                    i += 1
                
                # 处理普通段落
                else:
                    p = doc.add_paragraph()
                    MarkdownToWordConverter._add_formatted_text(p, line)
                    i += 1
        
        doc.save(output_path)
        print(f"✓ Word文档已保存至: {output_path}")
    
    @staticmethod
    def _add_formatted_text(paragraph, text):
        """处理文本中的粗体标记 **text**"""
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 粗体文本
                run = paragraph.add_run(part[2:-2])
                MarkdownToWordConverter.set_font(run, '微软雅黑', 11, True)
            elif part:
                # 普通文本
                run = paragraph.add_run(part)
                MarkdownToWordConverter.set_font(run, '微软雅黑', 11, False)



class MemoGenerator:
    """步骤3: 生成纪要"""
    def __init__(self, gemini_handler):
        self.gemini = gemini_handler
        self.memo_prompt_template = """ 
### 🎯 **配置区**


*   **目标发言人ID**：{speaker_id}


---


你是一名顶级的速记整理专家，任务是根据下方的逐字稿，以100%的颗粒度精准还原指定发言人的内容。你的核心原则是：**只做减法，不做任何改写或总结。** 目标是产出一份既保留了全部核心信息、又剔除了所有口语噪音的、高度可读的访谈纪要。


### **思维链处理步骤 (Chain-of-Thought Process)**


请严格遵循以下三个步骤，一步一步地处理用户提供的逐字稿：


**步骤 1: 识别与整合 (Identify & Integrate)**
*   **核心任务**: 通读全文，定位和识别出所有属于 `{speaker_id}` 的直接发言段落。
*   **整合规则**:
    1.  **整合确认信息**: 当 `{speaker_id}` 对另一位发言者的**事实性提问**做出简短的肯定性回答时（如："对"、"是的"、"没错"），应将提问中的**核心事实陈述**转化为一句陈述句，并视为 `{speaker_id}` 的发言。
    2.  **整合上下文信息**: 当另一位发言者的提问为 `{speaker_id}` 的回答提供了**理解所必需的关键上下文**（如具体的产品、价格范围、时间段等）时，应将该上下文信息自然地融入到回答，使回答本身成为一个独立完整的信息点。
*   **排除规则**: 完全忽略所有其他发言人的独立对话。进入下一步处理的，必须是 `{speaker_id}` 的直接发言和被他/她确认或补充了上下文的信息。


**步骤 2: 深度清理 (Deep Clean)**
*   然后，针对上一步识别和整合出的每一个段落，进行逐句处理。
*   在这一步，你的目标是**深度清理**，在**绝对不改变原文核心词汇和句子结构**的前提下，删除以下几类**不承载核心信息**的口语化表达：
    *   **基础语气词与口头禅**: “嗯”、“啊”、“呃”、“嘛”、“这个”、“那个”、“然后”、“就是”、“对”、“是的”等（注意：在步骤1中用于判断“确认”的“对/是的”在整合后应被删除）。
    *   **冗余的过渡与引言短语**: “你可以理解为”、“我的意思是”、“我想说的是”、“其实”、“实际上”、“怎么说呢”、“那我给你介绍一下”、“首先我先说一下”、“换句话说”、“简单来说”等。
    *   **犹豫与自我修正**: “我想一下”、“让我想想”、“说白了就是”、“不对，应该是”等。
*   **严格禁止**：进行任何主观改写、总结、增删或调整句子顺序。


**步骤 3: 结构化输出 (Structure & Format)**
*   **核心原则：零修改**。**绝对禁止**对原文进行任何文字上的增删、改写或总结。你**唯一能做**的是在不改变任何一个字的前提下，遵循以下规则进行格式化。
    *   **唯一例外**: 如果原文中出现了可根据上下文明确判断的、由语音转录造成的明显错别字（如同音/近音词错误、专有名词错误等），**允许且仅允许**进行修正。
*   **格式化规则 (Formatting Rules)**:
    1.  **输出格式**: 你的最终输出**只允许**包含以下 Markdown 元素：
        *   要点符号 (`-` 或 `*`)，包括缩进的二级要点。
        *   加粗 (`**文字**`)。
    2.  **禁止项**: **绝对禁止**在输出中使用任何级别的标题符号 (`#`)。
    3.  **布局**: 保持内容紧凑，不要添加不必要的空行。


---


### **🔹 示例说明**


*（请注意：以下示例旨在演示当您将 `[指定发言人]` 设置为“说话人2”时的预期输出效果。）*


**示例 1：基础清理**
原始逐字稿
说话人2：市场趋势方面，我们可以看到，电动汽车的销量增长很快。嗯，对，所以这也是为什么各大车企都在加速电动化布局。这个布局，我想，不仅仅是推出新车型，还包括了充电设施、电池技术等一系列的投入。


优化后输出
**市场趋势与车企布局**
- 市场趋势方面，我们可以看到，电动汽车的销量增长很快，这也是为什么各大车企都在加速电动化布局。
- 这个布局，不仅仅是推出新车型，还包括了：
  - 充电设施
  - 电池技术等一系列的投入。


**示例 2：整合确认信息**
原始逐字稿
说话人 1：嗯，那 1, 000 元每分钟差不多就是单部 10 万左右的成本，是吧？制作成本。
说话人 2：对， 10 ~ 20 万左右的一个水平，看时长的。
说话人 1：那这个 10 ~ 20 万是包括了从版权到最后完片成片，是吧？
说话人 2：对对对。


优化后输出
**单部剧制作成本**
- 1,000 元每分钟差不多是单部剧 10-20 万左右的制作成本，看时长，这包括了从版权到最后完片成片。


**示例 3：整合上下文与深度清理**
原始逐字稿
说话人 1：那比如说您说的那个 3, 000 到 4, 000 的这种有代表的作品吗？
说话人 2：我们现在在做的，比如说那个像那个叫什么？我不知道现在什么情况，那个叫什么狐？九尾狐，然后还有就是我九尾狐好像是上了，应该是已经上了这个项目没有跟，然后主要是这个，然后另外的话我们可能目前在做这一块的比较多，就是还在制作的过程中比较多。


优化后输出
**代表作品案例**
- 3,000 到 4, 000 的代表作品，我们现在在做的是《九尾狐》。它好像是上了，应该是已经上了，这个项目没有跟。主要是这个。另外，我们目前在做这一块的比较多，还在制作的过程中比较多。


---


### **📝 输出要求**
你的回复**必须**直接以 Markdown 格式的正文开始。**严禁**包含任何标题、开场白、解释、注释或任何非正文内容（例如，不要说“好的，这是整理好的内容：”）。
"""
    
    def _process_single_segment(self, seg_index, seg, speaker_id):
        """处理单个段落的辅助方法"""
        memo_prompt = self.memo_prompt_template.format(speaker_id=speaker_id)
        user_prompt = f"请处理以下逐字稿片段：\n\n{seg['content']}"
        
        result = self.gemini.request_gpt(
            system_prompt=memo_prompt,
            user_prompt=user_prompt,
            output_type="text",
            temperature=0
        )
        
        return {
            'index': seg_index,
            'topic': seg['topic'],
            'success': result['code'] == 0,
            'content': result['message'].strip() if result['code'] == 0 else None
        }
    
    def generate_memo(self, merged_segments, output_path, speaker_id):
        """并行处理所有段落生成最终纪要MD文档和Word文档"""
        print("=" * 60)
        print("步骤 3/3: 正在生成访谈纪要...")
        print("=" * 60 + "\n")
        
        print(f"目标说话人: {speaker_id}")
        print(f"共 {len(merged_segments)} 个段落待处理")
        print(f"✓ 启动并行处理模式\n")
        
        results = {}
        failed_segments = []
        
        with ThreadPoolExecutor() as executor:
            future_to_index = {
                executor.submit(self._process_single_segment, i, seg, speaker_id): i 
                for i, seg in enumerate(merged_segments, start=1)
            }
            
            for future in as_completed(future_to_index):
                try:
                    result = future.result()
                    results[result['index']] = result
                    
                    if result['success']:
                        print(f"✓ 段落 {result['index']} 处理成功: {result['topic']}")
                    else:
                        failed_segments.append(result['index'])
                        print(f"✗ 段落 {result['index']} 处理失败: {result['topic']}")
                        
                except Exception as e:
                    index = future_to_index[future]
                    failed_segments.append(index)
                    print(f"✗ 段落 {index} 处理异常: {str(e)}")
        
        markdown_content = ""
        for i in range(1, len(merged_segments) + 1):
            result = results.get(i)
            if result and result['success']:
                markdown_content += f"### {result['topic']}\n"
                markdown_content += result['content'] + "\n\n"
            else:
                topic = merged_segments[i-1]['topic']
                markdown_content += f"### {topic}\n[段落 {i} 处理失败]\n\n"
        
        # 保存MD文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        # 生成Word文档
        word_output_path = Path(output_path).with_suffix('.docx')
        MarkdownToWordConverter.convert(markdown_content, str(word_output_path))
        
        print(f"\n{'='*60}")
        print(f"✓ 纪要生成完成！")
        print(f"✓ MD文档已保存至: {output_path}")
        print(f"✓ Word文档已保存至: {word_output_path}")
        print(f"✓ 成功率: {len(merged_segments) - len(failed_segments)}/{len(merged_segments)}")
        print(f"{'='*60}\n")
        
        return markdown_content



if __name__ == "__main__":
    # 配置文件路径 - 现在同时支持 .txt 和 .docx，也支持文件夹批量处理
    TRANSCRIPT_PATH = r"E:\call summarizer\transcript\Interview recording 183.txt"  # 可以是文件或文件夹
    OUTPUT_DIR = Path(r"E:\call summarizer\output")  # 输出文件夹
    
    # 前端可指定目标说话人（留空则使用AI识别的专家）
    TARGET_SPEAKER = ""  # 例如: "说话人 2" 或留空 ""
    
    # 创建output文件夹（如果不存在）
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    print("\n" + "🚀 开始处理专家访谈..." + "\n")
    
    gemini = GeminiHandler()
    
    # 检查输入路径是文件还是文件夹
    input_path = Path(TRANSCRIPT_PATH)
    
    if input_path.is_file():
        # 单文件处理
        file_list = [input_path]
    elif input_path.is_dir():
        # 文件夹处理：收集所有txt和docx文件
        file_list = list(input_path.glob("*.txt")) + list(input_path.glob("*.docx"))
        if not file_list:
            print("\n❌ 文件夹中没有找到txt或docx文件")
            exit(1)
        print(f"📁 找到 {len(file_list)} 个文件待处理\n")
    else:
        print("\n❌ 输入路径无效")
        exit(1)
    
    # 循环处理每个文件
    for idx, transcript_file in enumerate(file_list, 1):
        if len(file_list) > 1:
            print(f"\n{'='*60}")
            print(f"处理文件 {idx}/{len(file_list)}: {transcript_file.name}")
            print(f"{'='*60}\n")
            
        # 生成输出文件名：原文件名_memo.md，统一输出到output文件夹
        output_file = OUTPUT_DIR / f"{transcript_file.stem}_memo.md"
        
        # 执行处理流程
        segmenter = TranscriptSegmenter(gemini)
        expert_speaker_id, segments_json = segmenter.segment_transcript(str(transcript_file))
        
        if not segments_json:
            print(f"\n❌ {transcript_file.name} 分段失败，跳过")
            continue
        
        # 确定使用的说话人ID：优先使用前端指定，否则使用AI识别的
        final_speaker_id = TARGET_SPEAKER if TARGET_SPEAKER else expert_speaker_id
        
        merger = TranscriptMerger()
        merged_segments = merger.merge_segments(segments_json, str(transcript_file))
        
        memo_gen = MemoGenerator(gemini)
        final_memo = memo_gen.generate_memo(merged_segments, str(output_file), final_speaker_id)
    
    print("\n🎉 全部流程完成！\n")
    if len(file_list) > 1:
        print(f"📄 共处理 {len(file_list)} 个文件")
    print(f"📂 输出目录: {OUTPUT_DIR}")